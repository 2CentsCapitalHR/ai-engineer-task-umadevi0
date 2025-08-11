#!/usr/bin/env python3
"""
Corporate Agent App — Document Review Tool for ADGM Compliance

Description:
    This script processes uploaded corporate documents (.docx) and checks them against
    ADGM regulatory requirements. It detects missing checklist items, flags potential
    compliance issues, and annotates documents with inline review comments.

Main Features:
    - Extracts text from DOCX and PDF files.
    - Matches detected document types to required ADGM checklists.
    - Detects 'red flag' issues such as missing jurisdiction clauses, absent signatures,
      placeholders, and ambiguous wording.
    - Uses a RAG (Retrieval-Augmented Generation) index for evidence retrieval.
    - Optionally integrates with LLM APIs (OpenAI / Hugging Face) to draft corrective clauses.
    - Annotates .docx files with inline highlights and generates a summary review page.
    - Provides a Gradio web interface for file upload, processing, and download.

Usage:
    python corporate_agent_app.py
"""
import os
import io
import json
import tempfile
import re
from typing import List, Dict, Any, Tuple, Optional

import gradio as gr
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

import PyPDF2

from sentence_transformers import SentenceTransformer, util as st_util
import numpy as np
import faiss

# LLM backends
try:
    from openai import OpenAI
except Exception:
    OpenAI = None
try:
    from huggingface_hub import InferenceClient
except Exception:
    InferenceClient = None

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# -----------------------------
# Base config (aligned with Data Sources.pdf)
# -----------------------------
CHECKLISTS = {
    "company_incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors",
        "Business Plan",
        "Signed Lease Agreement",
        "Passport Copies of Authorised Signatories",
        "Resolution for Incorporation",
    ],
    "branch_registration": [
        "Business Plan",
        "Parent Company Articles of Association",
        "Latest Audited Financial Accounts",
        "Resolution from Parent Company",
        "Signed Lease Agreement",
        "Passport Copies of Authorised Signatories",
        "Ownership Structure Chart",
    ],
    "data_protection": [
        "Appropriate Policy Document Template",
    ],
}

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Incorporation Application Form": ["incorporation application", "application form"],
    "UBO Declaration Form": ["ubo", "ultimate beneficial owner"],
    "Register of Members and Directors": ["register of members", "register of directors"],
    "Resolution for Incorporation": ["resolution for incorporation", "company formation resolution"],
    "Business Plan": ["business plan", "overview of the applicant", "target markets"],
    "Signed Lease Agreement": ["lease agreement", "registered office"],
    "Passport Copies of Authorised Signatories": ["authorised signatories", "passport copies"],
    "Appropriate Policy Document Template": ["appropriate policy document", "data protection"],
    "Parent Company Articles of Association": ["parent company articles", "articles of association"],
    "Latest Audited Financial Accounts": ["audited financial accounts", "financial statements"],
    "Resolution from Parent Company": ["resolution", "parent company resolution"],
    "Ownership Structure Chart": ["ownership structure", "ultimate beneficial owners"],
}

PROCESS_NAMES = {
    "Parent Company Articles of Association": "branch_registration",
    "Latest Audited Financial Accounts": "branch_registration",
    "Resolution from Parent Company": "branch_registration",
    "Ownership Structure Chart": "branch_registration",
    "Articles of Association": "company_incorporation",
    "Memorandum of Association": "company_incorporation",
    "Incorporation Application Form": "company_incorporation",
    "UBO Declaration Form": "company_incorporation",
    "Register of Members and Directors": "company_incorporation",
    "Resolution for Incorporation": "company_incorporation",
    "Business Plan": "company_incorporation",
    "Signed Lease Agreement": "company_incorporation",
    "Passport Copies of Authorised Signatories": "company_incorporation",
    "Appropriate Policy Document Template": "data_protection",
}

# -----------------------------
# Red flag rules: expanded
# -----------------------------
RED_FLAG_RULES = [
    {
        "id": "jurisdiction_wrong",
        "description": "Jurisdiction references non-ADGM court",
        "pattern": r"(UAE Federal Courts|Dubai Courts|Sharjah Courts|Abu Dhabi Courts|DIFC Courts|Sharjah Sharqiya Courts)",
        "severity": "High",
        "suggestion": "Replace jurisdiction references with ADGM Courts/ADGM jurisdiction per ADGM regulations.",
        "citation": "ADGM Companies Regulations"
    },
    {
        "id": "missing_signatory",
        "description": "Missing signatory or signature block",
        "pattern": r"(Signed|Signature|For and on behalf|authorised signatory|Signed by|signature block|Signed\s*:)",
        "severity": "Medium",
        "suggestion": "Add a signatory block including name, title, and date.",
        "citation": "Rule 12, Company Incorporation Package"
    },
    {
        "id": "missing_retention_policy",
        "description": "Missing or indefinite data retention policy",
        "pattern": r"(retain.*indefinitely|no retention policy|indefinite retention|retain for an indefinite period)",
        "severity": "High",
        "suggestion": "Specify a clear retention and erasure policy per DPR 2021 Article 7(3).",
        "citation": "DPR 2021 Article 7(3)"
    },
    {
        "id": "placeholder_fields",
        "description": "Placeholder fields or empty tables detected",
        "pattern": r"(\[\s*\]|\[___+\]|\_{6,}|<placeholder>|<name>|TO_BE_FILLED)",
        "severity": "High",
        "suggestion": "Fill in the placeholder fields (names, shares, amounts).",
        "citation": "ADGM Companies Regulations"
    },
]

# ambiguous / non-binding wording detector patterns
AMBIGUOUS_PATTERNS = [
    r"\bmay\b",
    r"\bmight\b",
    r"\bcould\b",
    r"\bwhere appropriate\b",
    r"\bas required\b",
    r"\breasonable efforts\b",
    r"\bsubject to\b",
    r"\bwhere possible\b",
    r"\bat its discretion\b",
    r"\bif practicable\b",
    r"\bto the extent practicable\b",
]

# -----------------------------
# Utilities: extract text
# -----------------------------
def extract_text_from_docx(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)

def extract_text_from_pdf(pdf_path: str) -> str:
    pages = []
    with open(pdf_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, p in enumerate(reader.pages):
            try:
                txt = p.extract_text() or ""
                #  Replace tabs with pipes for table parsing
                txt = txt.replace('\t', ' | ')
            except Exception:
                txt = ""
            pages.append(f"[PAGE {i+1}] {txt}")
    return "\n".join(pages)

# -----------------------------
# RAG index wrapper
# -----------------------------
class RAGIndex:
    def __init__(self, embed_model_name: str = "all-MiniLM-L6-v2"):
        self.embed_model = SentenceTransformer(embed_model_name)
        self.texts: List[str] = []
        self.index = None
        self.embeddings = None

    def build(self, docs: List[str]):
        chunks = []
        for d in docs:
            parts = re.split(r"\n{2,}|(?<=\.)\s+", d)
            for p in parts:
                s = p.strip()
                if len(s) >= 30:
                    chunks.append(s)
        self.texts = chunks
        if not chunks:
            return
        embs = self.embed_model.encode(self.texts, convert_to_numpy=True)
        self.embeddings = embs
        d = embs.shape[1]
        self.index = faiss.IndexFlatL2(d)
        self.index.add(embs)

    def query(self, q: str, top_k: int = 3) -> List[Tuple[str, float]]:
        if self.index is None or not self.texts:
            return []
        q_emb = self.embed_model.encode([q], convert_to_numpy=True)
        D, I = self.index.search(q_emb, top_k)
        res = []
        for idx, dist in zip(I[0], D[0]):
            if idx < 0 or idx >= len(self.texts):
                continue
            #  convert faiss L2 distance to a pseudo-score (lower = closer)
            res.append((self.texts[idx], float(dist)))
        return res

# -----------------------------
# Extract checklist items from datasources
# -----------------------------
def extract_checklist_items_from_datasources(texts: List[str]) -> List[str]:
    items = []
    for t in texts:
        lines = [l.strip() for l in t.splitlines() if l.strip()]
        i = 0
        current_category = ""
        current_doc_type = []
        while i < len(lines):
            line = lines[i]
            i += 1
            #  Skip headers, URLs, and notes
            if 'Category' in line or 'Document/Template Type' in line or 'Official ADGM/Government Link' in line or re.match(r'https?://', line) or line.startswith('Note:'):
                continue
            #  Handle category
            if re.match(r'^[A-Z][A-Za-z\s&]{5,50}$', line) and not current_doc_type:
                current_category = line.strip()
                continue
            #  Handle document type
            if not re.match(r'https?://', line):
                current_doc_type.append(line)
            #  Process when a new category, URL, or end is reached
            if i >= len(lines) or re.match(r'^[A-Z][A-Za-z\s&]{5,50}$', lines[i]) or re.match(r'https?://', lines[i]):
                if current_doc_type:
                    doc_type_str = ' '.join(current_doc_type).strip()
                    if doc_type_str and not re.match(r'https?://', doc_type_str):
                        full = f"{current_category} {doc_type_str}".strip() if current_category else doc_type_str
                        items.append(full)
                    current_doc_type = []
        #  Save last doc_type
        if current_doc_type:
            doc_type_str = ' '.join(current_doc_type).strip()
            if doc_type_str and not re.match(r'https?://', doc_type_str):
                full = f"{current_category} {doc_type_str}".strip() if current_category else doc_type_str
                items.append(full)
    #  Clean and deduplicate
    cleaned = []
    seen = set()
    for it in items:
        it2 = re.sub(r'[^A-Za-z0-9 &\-/(),–+]', '', it).strip()
        if len(it2) < 10 or 'http' in it2.lower() or 'note' in it2.lower():
            continue
        #  Match to known keywords
        for doc_type, keywords in DOC_TYPE_KEYWORDS.items():
            if any(kw in it2.lower() for kw in keywords):
                it2 = doc_type
                break
        if it2.lower() not in seen:
            cleaned.append(it2)
            seen.add(it2.lower())
    return cleaned or CHECKLISTS.get("branch_registration", [])

# -----------------------------
# Document analysis helpers
# -----------------------------
def detect_document_types(text: str) -> List[str]:
    found = set()
    tl = text.lower()
    for dtype, keywords in DOC_TYPE_KEYWORDS.items():
        for kw in keywords:
            if kw in tl:
                found.add(dtype)
    return list(found)

def compare_against_checklist(detected: List[str], datasource_items: List[str], process: str) -> Dict[str, Any]:
    #  Use static checklist for the detected process
    required = CHECKLISTS.get(process, CHECKLISTS["company_incorporation"])
    present = []
    for r in required:
        for d in detected:
            if r.lower() in d.lower() or d.lower() in r.lower():
                present.append(r)
                break
    missing = [r for r in required if r not in present]
    return {
        "process": process,
        "documents_uploaded": len(detected),
        "required_documents": len(required),
        "missing_documents": missing
    }

def find_red_flags(text: str, rag: Optional[RAGIndex] = None, top_k_evidence: int = 3) -> List[Dict[str, Any]]:
    """
    Detect red flags using rules + ambiguous language detector.
    Enrich each issue with RAG evidence passages (if RAG provided).
    """
    issues = []
    tl = text or ""
    #  Rule-based flags
    for rule in RED_FLAG_RULES:
        if re.search(rule["pattern"], tl, flags=re.IGNORECASE):
            evidence = []
            if rag and rag.index is not None:
                passages = rag.query(rule["description"], top_k=top_k_evidence)
                for p, dist in passages:
                    evidence.append({"source_snippet": p, "distance": dist})
            issues.append({
                "document": None,
                "section": None,
                "issue": rule["description"],
                "rule_id": rule["id"],
                "severity": rule["severity"],
                "suggestion": rule["suggestion"],
                "citation": rule.get("citation", "ADGM Regulations"),
                "evidence": evidence
            })
    #  Signature detection : look for typical signature blocks or absence thereof
    sig_patterns = [
        r"signature\s*[:\-]", r"signed\s*by", r"for and on behalf of", r"authorised signatory", r"signatory\s*name",
        r"__________________", r"____________________"
    ]
    if not any(re.search(p, tl, flags=re.IGNORECASE) for p in sig_patterns):
        issues.append({
            "document": None,
            "section": "End",
            "issue": "No explicit signature block found",
            "rule_id": "missing_signatory_block",
            "severity": "Medium",
            "suggestion": "Include an explicit signature block (name, title, date).",
            "citation": "Rule 12, Company Incorporation Package",
            "evidence": []
        })
    #  Placeholder fields / blanks (already covered in RED_FLAG_RULES but double-check)
    if re.search(r'(\[\s*\]|\[___+\]|\_{6,}|<placeholder>|TO_BE_FILLED)', tl, flags=re.IGNORECASE):
        issues.append({
            "document": None,
            "section": None,
            "issue": "Placeholder fields or empty tables detected",
            "rule_id": "placeholder_fields",
            "severity": "High",
            "suggestion": "Fill in the placeholder fields (names, shares, amounts).",
            "citation": "ADGM Companies Regulations",
            "evidence": []
        })
    #  Ambiguous / non-binding language detection
    ambiguous_matches = []
    for pat in AMBIGUOUS_PATTERNS:
        for m in re.finditer(pat, tl, flags=re.IGNORECASE):
            snippet = tl[max(0, m.start()-60): m.end()+60].strip()
            ambiguous_matches.append(snippet)
    if ambiguous_matches:
        #  collect a few unique snippets
        unique_snips = []
        for s in ambiguous_matches:
            if s not in unique_snips:
                unique_snips.append(s)
            if len(unique_snips) >= 5:
                break
        evidence = []
        if rag and rag.index is not None:
            #  ask RAG for relevant ADGM normative text for "binding language" guidance
            passages = rag.query("binding language / mandatory wording best practice", top_k=2)
            for p, dist in passages:
                evidence.append({"source_snippet": p, "distance": dist})
        issues.append({
            "document": None,
            "section": None,
            "issue": "Ambiguous or non-binding language detected",
            "rule_id": "ambiguous_language",
            "severity": "Medium",
            "suggestion": "Replace ambiguous terms (e.g., 'may', 'reasonable efforts') with clear, mandatory obligations where appropriate.",
            "citation": "ADGM drafting guidance (see ADGM Companies Regulations for specific clauses)",
            "evidence": evidence,
            "occurrences": unique_snips
        })
    return issues

# -----------------------------
# Annotate docx: inline + summary page 
# -----------------------------
def _shorten(text: str, n: int = 300) -> str:
    if not text:
        return ""
    return (text[:n] + "...") if len(text) > n else text

def annotate_docx(docx_bytes: bytes, findings: List[Dict[str, Any]]) -> bytes:
    """
    Insert short inline annotations noting issue + citation + suggested clause snippet.
    Also appends an 'Automated Review Summary' page with full details.
    """
    doc = Document(io.BytesIO(docx_bytes))

    #  Build a map of patterns to annotations to avoid repeating many times
    #  We'll search for key words from the issue description or specific occurrences
    for p in doc.paragraphs:
        text = p.text or ""
        low = text.lower()
        added_any = False
        for f in findings:
            issue_text = f.get("issue", "")
            #  if issue has specific 'occurrences' list use that, else use first 6 words of issue
            occurrences = f.get("occurrences") or []
            matched = False
            #  Try explicit occurrence snippets first
            for occ in occurrences:
                if occ and occ.lower().strip()[:10] in low:
                    matched = True
                    break
            #  fallback: match by key words from issue_text
            if not matched:
                key = " ".join(issue_text.split()[:6]).lower()
                if key and key in low:
                    matched = True
            if matched:
                #  Add an inline annotation run (yellow highlight) with short suggestion + citation
                sug = f.get("suggestion", "See review.")
                citation = f.get("citation", "")
                suggested_clause = f.get("suggested_clause")
                #  Shorten suggested clause to keep inline concise
                short_clause = _shorten(str(suggested_clause or ""), 300)
                ann_text_parts = [f"[REVIEW: {sug}"]
                if citation:
                    ann_text_parts.append(f"| Citation: {citation}")
                if short_clause:
                    ann_text_parts.append(f"| Suggested: {short_clause}")
                ann_text = " ".join(ann_text_parts) + "]"
                try:
                    run = p.add_run(" " + ann_text)
                    run.font.size = Pt(9)
                    try:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except Exception:
                        pass
                    added_any = True
                except Exception:
                    #  if paragraph run fails, append as new paragraph
                    np_ = doc.add_paragraph(ann_text)
                    r = np_.runs[0] if np_.runs else None
                    if r:
                        r.font.size = Pt(9)
                        try:
                            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        except Exception:
                            pass
        

    #  Append summary page
    doc.add_page_break()
    try:
        doc.add_heading('Automated Review Summary', level=2)
    except KeyError:
        p = doc.add_paragraph()
        r = p.add_run("Automated Review Summary")
        r.bold = True
        p.style = 'Normal'

    for idx, f in enumerate(findings, start=1):
        doc.add_paragraph(f"{idx}. Issue: {f.get('issue')} | Severity: {f.get('severity')} | Citation: {f.get('citation')}")
        if f.get('suggestion'):
            doc.add_paragraph(f"   Suggestion: {f.get('suggestion')}")
        if f.get('suggested_clause'):
            doc.add_paragraph(f"   Suggested clause (LLM): {str(f.get('suggested_clause'))[:2000]}")
        #  Evidence
        evs = f.get('evidence') or []
        for ev in evs[:4]:
            src = ev.get('source_snippet') or ""
            doc.add_paragraph(f"   Evidence: {src[:800]}")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# -----------------------------
# LLM wrapper (Hugging Face or OpenAI)
# -----------------------------
def call_llm(prompt: str, api_key: Optional[str] = None, model: str = "mistralai/Mixtral-8x7B-Instruct-v0.1") -> str:
    """
    If api_key starts with 'hf_' -> use Hugging Face InferenceClient.
    Else assume OpenAI key. Returns string or an error string.
    """
    if not api_key:
        return "(LLM skipped — no API key provided)"

    api_key = api_key.strip()
    #  Hugging Face token path
    if api_key.startswith("hf_") and InferenceClient:
        try:
            client = InferenceClient(token=api_key, model=model)
            out = client.text_generation(prompt=prompt, max_new_tokens=200)
            if isinstance(out, (list, tuple)) and len(out):
                return out[0].get("generated_text", str(out[0]))
            if isinstance(out, dict):
                return out.get("generated_text", str(out))
            return str(out)
        except Exception as e:
            return f"(LLM skipped — Hugging Face error: {e})"

    #  OpenAI path
    if OpenAI:
        try:
            client = OpenAI(api_key=api_key)
            resp = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=300,
                temperature=0.0,
            )
            return resp.choices[0].message.content
        except Exception as e:
            return f"(LLM skipped — OpenAI error: {e})"

    return "(LLM skipped — no LLM backend available)"

# -----------------------------
# Main pipeline
# -----------------------------



def process_documents(files: List[Tuple[str, bytes]], datasources_paths: List[str], api_key: Optional[str] = None) -> Dict[str, Any]:
    results: Dict[str, Any] = {
        "documents": [],
        "process": None,
        "documents_uploaded": 0,
        "required_documents": 0,
        "missing_documents": [],
        "issues_found": [],
        "datasource_items": [],
        "notification": "",
        "satisfied_items_map": {},
    }

    #  Read datasources
    datasource_texts = []
    for p in datasources_paths or []:
        if p.lower().endswith(".pdf"):
            try:
                datasource_texts.append(extract_text_from_pdf(p))
            except Exception as e:
                print(f"Failed to extract pdf {p}: {e}")
        else:
            try:
                with open(p, "r", encoding="utf-8") as f:
                    datasource_texts.append(f.read())
            except Exception:
                pass

    #  Build RAG
    rag = RAGIndex()
    if datasource_texts:
        rag.build(datasource_texts)

    #  Extract checklist items from datasource files
    datasource_items = extract_checklist_items_from_datasources(datasource_texts)

    #  Determine process
    process = None
    detected_types_all = []
    for fname, fbytes in files:
        try:
            text = extract_text_from_docx(fbytes)
            dtypes = detect_document_types(text)
            detected_types_all.extend(dtypes)
            for dtype in dtypes:
                if dtype in PROCESS_NAMES:
                    process = PROCESS_NAMES[dtype]
                    if process == "branch_registration":  # Prioritize branch_registration
                        break
        except Exception:
            continue
    if not process:
        process = "company_incorporation"

    #  using the static checklist for the detected process 
    static_required = CHECKLISTS.get(process, CHECKLISTS["company_incorporation"])
    if not datasource_items:
        datasource_items = static_required
    results["datasource_items"] = datasource_items

    #  Precompute embeddings for fuzzy matching
    embed_model = rag.embed_model if rag and rag.embed_model else SentenceTransformer("all-MiniLM-L6-v2")
    try:
        datasource_item_embs = embed_model.encode(static_required, convert_to_numpy=True) if static_required else np.array([])  
    except Exception:
        datasource_item_embs = np.array([])

    satisfied_items_map = {}

    #  Process each uploaded doc
    for fname, fbytes in files:
        entry = {
            "filename": fname,
            "detected_types": [],
            "issues": []
        }
        try:
            text = extract_text_from_docx(fbytes)
        except Exception as e:
            entry["issues"].append({
                "document": fname,
                "issue": "Invalid or corrupted .docx file",
                "severity": "Error",
                "suggestion": "Upload a valid .docx file."
            })
            results["documents"].append(entry)
            continue

        #  Detect types
        dtypes = detect_document_types(text)
        entry["detected_types"] = dtypes

        #  Red flags
        rflags = find_red_flags(text, rag=rag)
        enriched_flags = []
        for rf in rflags:
            evidence = rf.get("evidence", [])
            if not evidence and rag and rag.index is not None:
                passages = rag.query(rf.get("issue", ""), top_k=3)
                for p, dist in passages:
                    evidence.append({"source_snippet": p, "distance": dist})
            ctx_text = "\n\n".join([ev["source_snippet"] for ev in evidence]) if evidence else ""
            prompt = (
                f"You are a legal assistant. Draft a concise corrective clause (1-3 sentences) "
                f"for the issue: '{rf.get('issue')}'.\n\nRelevant ADGM references:\n{ctx_text}\n\nDraft the clause in plain language and keep it brief."
            )
            suggested = call_llm(prompt, api_key=api_key) if api_key else "(LLM skipped - no API key)"
            rf["suggested_clause"] = suggested
            rf["evidence"] = evidence
            enriched_flags.append(rf)
        entry["issues"].extend(enriched_flags)

        #  Fuzzy-match to static checklist 
        satisfied = []
        if static_required and datasource_item_embs.size and embed_model:
            try:
                #  Exact matches from detected_types
                for dtype in dtypes:
                    if dtype in static_required and not any(s["item"] == dtype for s in satisfied):
                        satisfied.append({"item": dtype, "score": 1.0})
                #  Fuzzy matches if not already satisfied
                doc_emb = embed_model.encode([text], convert_to_numpy=True)
                sims = st_util.cos_sim(doc_emb, datasource_item_embs)[0].cpu().numpy()
                THRESH = 0.7
                for idx, score in enumerate(sims):
                    if float(score) >= THRESH:
                        if not any(s["item"] == static_required[idx] for s in satisfied):
                            satisfied.append({"item": static_required[idx], "score": float(score)})
            except Exception:
                pass
        satisfied_items_map[fname] = satisfied

        results["documents"].append(entry)

    #  Collate covered items against static checklist 
    covered = set()
    for v in satisfied_items_map.values():
        for s in v:
            covered.add(s["item"].lower())

    missing = [it for it in static_required if it.lower() not in covered]

    results.update({
        "process": process,
        "documents_uploaded": len(files),
        "required_documents": len(static_required),
        "missing_documents": missing,
        "satisfied_items_map": satisfied_items_map,
    })

    #  Notification
    if missing:
        results["notification"] = (
            f"It appears that you’re trying to {process.lower()} in ADGM. "
            f"Based on our reference list, you have uploaded {len(files)} out of {len(static_required)} required documents. "
            f"The missing document(s) appear(s) to be: {', '.join(missing)}."
        )

    #  Collate issues and annotate
    all_issues = []
    for d in results["documents"]:
        for issue in d.get("issues", []):
            issue["document"] = d["filename"]
            all_issues.append(issue)
    results["issues_found"] = all_issues

    annotated_files = []
    for fname, fbytes in files:
        annotated_bytes = annotate_docx(fbytes, all_issues)
        p = os.path.join(tempfile.gettempdir(), f"reviewed_{fname}")
        with open(p, "wb") as fh:
            fh.write(annotated_bytes)
        annotated_files.append(p)

    results["annotated_files"] = annotated_files

    return results



# -----------------------------
# Gradio UI
# -----------------------------
def gradio_interface():
    with gr.Blocks() as demo:
        gr.Markdown("# ADGM Corporate Agent — Document Intelligence (Demo) — Updated")
        with gr.Row():
            doc_input = gr.File(label="Upload .docx files", file_count="multiple", file_types=['.docx'])
            ds_input = gr.File(label="Upload ADGM Data Sources (PDF/Text) — REQUIRED", file_count="multiple")
        api_key = gr.Textbox(label="LLM API Key (Hugging Face hf_... OR OpenAI sk_... — optional)")
        run_btn = gr.Button("Run Review")

        output_json = gr.JSON(label="Review JSON Output")
        download_files = gr.Files(label="Reviewed Files (download)")

        def run_review(docs, datasources, api_key):
            api_key_str = str(api_key).strip() if api_key and str(api_key).strip() else None
            uploaded = []
            for d in docs or []:
                fn = os.path.basename(d.name)
                with open(d.name, "rb") as fh:
                    uploaded.append((fn, fh.read()))
            ds_paths = [d.name for d in (datasources or [])]
            res = process_documents(uploaded, ds_paths, api_key_str)
            return res, res.get("annotated_files", [])

        run_btn.click(fn=run_review, inputs=[doc_input, ds_input, api_key], outputs=[output_json, download_files])

    return demo

if __name__ == "__main__":
    app = gradio_interface()
    app.launch(server_name='127.0.0.1', share=True)
